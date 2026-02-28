"""Word document (.docx) reader tool via python-docx.

Privilege level 1 — safe read-only document parsing.
"""

from typing import Any

import httpx
from pydantic import Field

from aumos_agent_framework.core.interfaces import ToolInputSchema, ToolOutputSchema
from aumos_common.observability import get_logger

logger = get_logger(__name__)


class DocxReaderInput(ToolInputSchema):
    """Input schema for the DOCX reader tool."""

    source_url: str = Field(..., description="URL to fetch the .docx file from")
    include_tables: bool = Field(default=True, description="Whether to include table cell text")
    max_chars: int = Field(
        default=50000,
        ge=1000,
        le=500000,
        description="Maximum characters to return",
    )


class DocxReaderOutput(ToolOutputSchema):
    """Output schema for the DOCX reader tool."""

    result: str = Field(description="Extracted text content from the Word document")
    metadata: dict[str, Any] = Field(default_factory=dict)


class DocxReaderTool:
    """Extract text from a Microsoft Word .docx file fetched from a URL.

    Uses python-docx for pure-Python DOCX parsing. Extracts paragraph text
    and optionally table cell text in reading order.
    """

    tool_id: str = "docx_reader"
    display_name: str = "Word Document Reader"
    category: str = "document"
    description: str = "Fetch a .docx file from a URL and extract its text content including paragraphs and tables."
    privilege_level: int = 1
    input_schema: type[DocxReaderInput] = DocxReaderInput
    output_schema: type[DocxReaderOutput] = DocxReaderOutput

    async def execute(
        self,
        input_data: DocxReaderInput,
        tenant_id: str,
        config: dict[str, str],
    ) -> DocxReaderOutput:
        """Fetch and extract text from a DOCX file.

        Args:
            input_data: DOCX URL, table inclusion flag, and char limit.
            tenant_id: Tenant context for logging.
            config: Optional 'AUTH_TOKEN' for authenticated document URLs.

        Returns:
            DocxReaderOutput with extracted text.
        """
        try:
            import io

            import docx
        except ImportError:
            return DocxReaderOutput(result="", metadata={"error": "python-docx not installed"})

        headers: dict[str, str] = {}
        auth_token = config.get("AUTH_TOKEN")
        if auth_token:
            headers["Authorization"] = f"Bearer {auth_token}"

        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.get(input_data.source_url, headers=headers)
            response.raise_for_status()

        document = docx.Document(io.BytesIO(response.content))
        text_parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        if input_data.include_tables:
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

        full_text = "\n".join(text_parts)[: input_data.max_chars]

        logger.info(
            "DOCX read",
            source_url=input_data.source_url,
            paragraphs=len(document.paragraphs),
            tables=len(document.tables),
            chars_extracted=len(full_text),
            tenant_id=tenant_id,
        )

        return DocxReaderOutput(
            result=full_text,
            metadata={
                "source_url": input_data.source_url,
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "chars_extracted": len(full_text),
            },
        )


TOOL = DocxReaderTool()
